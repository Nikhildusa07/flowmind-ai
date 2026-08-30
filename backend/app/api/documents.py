import csv
import io
import os
from pathlib import Path

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile, status
from sqlalchemy.orm import Session

from app.core.database import get_db
from app.models.models import Document, DocumentStatus, User
from app.services.ai_service import generate_summary
from app.core.dependencies import get_current_user

# IMPORTANT:
# If your project uses a different location for get_current_user,
# keep the dependency/import that your existing auth routes use.


router = APIRouter(
    prefix="/documents",
    tags=["Documents"],
)


# ============================================================
# CONFIGURATION
# ============================================================

ALLOWED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".csv",
    ".xlsx",
    ".xls",
}

# Keep this consistent with the frontend.
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 MB

UPLOAD_DIRECTORY = Path("uploads/documents")


# ============================================================
# FILE HELPERS
# ============================================================


def get_file_extension(filename: str) -> str:
    return os.path.splitext(filename)[1].lower()


def get_safe_filename(filename: str) -> str:
    """
    Prevent path traversal and keep only the filename.
    """
    return os.path.basename(filename)


def save_uploaded_file(
    content: bytes,
    filename: str,
) -> str:
    """
    Save the uploaded document locally and return its path.
    """

    UPLOAD_DIRECTORY.mkdir(
        parents=True,
        exist_ok=True,
    )

    safe_filename = get_safe_filename(filename)

    file_path = UPLOAD_DIRECTORY / safe_filename

    # Avoid overwriting an existing file with the same name.
    if file_path.exists():
        stem = file_path.stem
        suffix = file_path.suffix

        counter = 1

        while file_path.exists():
            file_path = (
                UPLOAD_DIRECTORY
                / f"{stem}_{counter}{suffix}"
            )

            counter += 1

    file_path.write_bytes(content)

    return str(file_path)


# ============================================================
# TEXT EXTRACTION
# ============================================================


def extract_txt(content: bytes) -> str:
    return content.decode(
        "utf-8",
        errors="ignore",
    )


def extract_csv(content: bytes) -> str:
    text = content.decode(
        "utf-8",
        errors="ignore",
    )

    reader = csv.reader(
        io.StringIO(text)
    )

    rows = []

    for row in reader:
        rows.append(
            " | ".join(
                str(value)
                for value in row
            )
        )

    return "\n".join(rows)


def extract_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(
            io.BytesIO(content)
        )

        pages = []

        for page in reader.pages:
            text = page.extract_text() or ""
            pages.append(text)

        return "\n".join(pages).strip()

    except ImportError:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="PDF processing dependency is not installed.",
        )


def extract_docx(content: bytes) -> str:
    try:
        from docx import Document as DocxDocument

        document = DocxDocument(
            io.BytesIO(content)
        )

        paragraphs = [
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]

        return "\n".join(paragraphs).strip()

    except ImportError:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="DOCX processing dependency is not installed.",
        )


def extract_excel(content: bytes) -> str:
    try:
        import pandas as pd

        dataframe = pd.read_excel(
            io.BytesIO(content)
        )

        return dataframe.to_string(
            index=False
        )

    except ImportError:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Spreadsheet processing dependency is not installed.",
        )


def extract_document_text(
    content: bytes,
    extension: str,
) -> str:

    if extension == ".txt":
        return extract_txt(content)

    if extension == ".csv":
        return extract_csv(content)

    if extension == ".pdf":
        return extract_pdf(content)

    if extension == ".docx":
        return extract_docx(content)

    if extension in {".xlsx", ".xls"}:
        return extract_excel(content)

    raise HTTPException(
        status_code=status.HTTP_400_BAD_REQUEST,
        detail="Unsupported document type.",
    )


# ============================================================
# DOCUMENT READING
# ============================================================


async def read_and_extract_document(
    file: UploadFile,
) -> tuple[str, str, str, bytes]:

    if not file.filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Filename is required.",
        )

    filename = get_safe_filename(
        file.filename
    )

    extension = get_file_extension(
        filename
    )

    if extension not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=(
                "Unsupported file type. Please upload "
                "PDF, DOCX, TXT, CSV, XLSX or XLS."
            ),
        )

    content = await file.read()

    if not content:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Uploaded document is empty.",
        )

    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail="Document size must not exceed 20 MB.",
        )

    try:
        extracted_text = extract_document_text(
            content,
            extension,
        )

    except HTTPException:
        raise

    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unable to process document: {str(exc)}",
        )

    if not extracted_text.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Document does not contain readable text.",
        )

    file_type = extension.replace(
        ".",
        "",
    ).upper()

    return (
        filename,
        file_type,
        extracted_text,
        content,
    )


# ============================================================
# AI PROCESSING
# ============================================================


def process_ai_summary(
    extracted_text: str,
) -> dict:

    try:
        result = generate_summary(
            extracted_text
        )

        if not result["success"]:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=result["message"],
            )

        return result

    except HTTPException:
        raise

    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"AI analysis failed: {str(exc)}",
        )


# ============================================================
# ANALYZE DOCUMENT
# ============================================================


@router.post(
    "/analyze",
    status_code=status.HTTP_200_OK,
)
async def analyze_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):

    filename = None
    saved_file_path = None

    try:
        (
            filename,
            file_type,
            extracted_text,
            content,
        ) = await read_and_extract_document(
            file
        )

        # --------------------------------------------------------
        # AI ANALYSIS
        # --------------------------------------------------------

        result = process_ai_summary(
            extracted_text
        )

        # --------------------------------------------------------
        # SAVE FILE
        # --------------------------------------------------------

        saved_file_path = save_uploaded_file(
            content,
            filename,
        )

        # --------------------------------------------------------
        # SAVE DOCUMENT RECORD
        # --------------------------------------------------------

        document = Document(
            filename=filename,
            file_type=file_type,
            file_path=saved_file_path,
            status=DocumentStatus.PROCESSED,
            extracted_text=extracted_text,
            summary=result["summary"],
            owner_id=current_user.id,
        )

        db.add(document)
        db.commit()
        db.refresh(document)

        # --------------------------------------------------------
        # RESPONSE
        # --------------------------------------------------------

        return {
            "success": True,
            "document_id": document.id,
            "filename": filename,
            "file_type": file_type,
            "summary": result["summary"],
            "insights": result["insights"],
            "status": document.status.value,
            "message": "Document summary generated successfully.",
        }

    except HTTPException:
        db.rollback()

        # If database saving happened to fail after the file
        # was saved, remove the local file.
        if saved_file_path:
            try:
                if os.path.exists(saved_file_path):
                    os.remove(saved_file_path)
            except Exception:
                pass

        raise

    except Exception as exc:
        db.rollback()

        if saved_file_path:
            try:
                if os.path.exists(saved_file_path):
                    os.remove(saved_file_path)
            except Exception:
                pass

        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Document analysis failed: {str(exc)}",
        )


# ============================================================
# SUMMARIZE DOCUMENT
# ============================================================


@router.post(
    "/summarize",
    status_code=status.HTTP_200_OK,
)
async def summarize_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):

    filename = None
    saved_file_path = None

    try:
        (
            filename,
            file_type,
            extracted_text,
            content,
        ) = await read_and_extract_document(
            file
        )

        result = process_ai_summary(
            extracted_text
        )

        saved_file_path = save_uploaded_file(
            content,
            filename,
        )

        document = Document(
            filename=filename,
            file_type=file_type,
            file_path=saved_file_path,
            status=DocumentStatus.PROCESSED,
            extracted_text=extracted_text,
            summary=result["summary"],
            owner_id=current_user.id,
        )

        db.add(document)
        db.commit()
        db.refresh(document)

        return {
            "success": True,
            "document_id": document.id,
            "filename": filename,
            "file_type": file_type,
            "summary": result["summary"],
            "insights": result["insights"],
            "status": document.status.value,
            "message": result["message"],
        }

    except HTTPException:
        db.rollback()

        if saved_file_path:
            try:
                if os.path.exists(saved_file_path):
                    os.remove(saved_file_path)
            except Exception:
                pass

        raise

    except Exception as exc:
        db.rollback()

        if saved_file_path:
            try:
                if os.path.exists(saved_file_path):
                    os.remove(saved_file_path)
            except Exception:
                pass

        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Document summarization failed: {str(exc)}",
        )